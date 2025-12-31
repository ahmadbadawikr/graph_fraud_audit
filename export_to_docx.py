"""
Export PROJECT_PAPER.md to DOCX with embedded images
Improved version - handles all content properly
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
MD_FILE = "/Users/kasyfur/graph_fraud_audit/PROJECT_PAPER.md"
OUTPUT_FILE = "/Users/kasyfur/graph_fraud_audit/PROJECT_PAPER.docx"
BASE_DIR = "/Users/kasyfur/graph_fraud_audit"

# Image path mapping
IMAGE_MAPPINGS = {
    'fig10_graph_schema.png': 'fig11_graph_schema.png',
    'fig11_hgt_architecture.png': 'fig12_hgt_architecture.png',
    'fig7_experiment_timeline.png': 'fig10_experiment_timeline.png',
    'fig1_training_curves.png': 'fig02_individual_training.png',
    'fig2_overfitting_analysis.png': 'fig08_v2_v3_comparison.png',
    'fig4_confusion_matrices.png': 'fig03_all_confusion_matrices.png',
    'fig7_feature_importance.png': 'fig05_complexity_performance.png',
    'fig6_ensemble_weights.png': 'fig09_ensemble_hybrid.png',
}

def clean_latex(text):
    """Convert LaTeX to readable text"""
    # Handle block equations
    text = re.sub(r'\$\$(.*?)\$\$', lambda m: f'[Equation: {clean_latex_symbols(m.group(1))}]', text, flags=re.DOTALL)
    # Handle inline math
    text = re.sub(r'\$([^\$]+?)\$', lambda m: clean_latex_symbols(m.group(1)), text)
    return text

def clean_latex_symbols(latex):
    """Clean LaTeX symbols to Unicode"""
    if not latex:
        return ''
    latex = latex.strip()
    
    # Arrows with text
    latex = re.sub(r'\\xrightarrow\{\\text\{([^}]*)\}\}', r' →[\1]→ ', latex)
    latex = re.sub(r'\\xrightarrow\{([^}]*)\}', r' →[\1]→ ', latex)
    
    # Subscripts and superscripts
    latex = re.sub(r'\^{\(([^}]*)\)}', r'^(\1)', latex)
    latex = re.sub(r'_\{([^}]*)\}', r'_\1', latex)
    latex = re.sub(r'\^\{([^}]*)\}', r'^(\1)', latex)
    
    # Greek letters
    replacements = {
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
        r'\epsilon': 'ε', r'\sigma': 'σ', r'\theta': 'θ', r'\lambda': 'λ',
        r'\mu': 'μ', r'\pi': 'π', r'\phi': 'φ', r'\psi': 'ψ', r'\omega': 'ω',
        r'\hat{y}': 'ŷ', r'\cdot': '·', r'\times': '×', r'\div': '÷',
        r'\pm': '±', r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
        r'\approx': '≈', r'\infty': '∞', r'\sum': 'Σ', r'\prod': 'Π',
        r'\rightarrow': '→', r'\leftarrow': '←', r'\Rightarrow': '⇒',
        r'\forall': '∀', r'\exists': '∃', r'\in': '∈',
    }
    for old, new in replacements.items():
        latex = latex.replace(old, new)
    
    # Clean up special functions
    latex = re.sub(r'\\mathcal\{([A-Z])\}', r'\1', latex)
    latex = re.sub(r'\\text\{([^}]*)\}', r'\1', latex)
    latex = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1/\2)', latex)
    latex = re.sub(r'\\[a-zA-Z]+', '', latex)
    latex = latex.replace('{', '').replace('}', '')
    
    return latex.strip()

def fix_image_path(path):
    """Fix image path using mapping"""
    filename = os.path.basename(path)
    if filename in IMAGE_MAPPINGS:
        dirname = os.path.dirname(path)
        return os.path.join(dirname, IMAGE_MAPPINGS[filename])
    return path

def parse_markdown(content):
    """Parse markdown into structured blocks"""
    content = clean_latex(content)
    lines = content.split('\n')
    blocks = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Code block
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            blocks.append({'type': 'code', 'content': '\n'.join(code_lines)})
            i += 1
            continue
        
        # Table - look for pipe characters
        if '|' in line and not line.startswith('!['):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            if table_lines:
                blocks.append({'type': 'table', 'lines': table_lines})
            continue
        
        # Heading
        if line.startswith('# '):
            blocks.append({'type': 'h1', 'content': line[2:].strip()})
        elif line.startswith('## '):
            blocks.append({'type': 'h2', 'content': line[3:].strip()})
        elif line.startswith('### '):
            blocks.append({'type': 'h3', 'content': line[4:].strip()})
        elif line.startswith('#### '):
            blocks.append({'type': 'h4', 'content': line[5:].strip()})
        
        # Image
        elif line.startswith('!['):
            match = re.search(r'\!\[.*?\]\((.*?)\)', line)
            if match:
                blocks.append({'type': 'image', 'path': match.group(1)})
        
        # Horizontal rule
        elif line.strip() == '---':
            blocks.append({'type': 'hr'})
        
        # Block quote
        elif line.startswith('> '):
            quote_lines = [line[2:]]
            while i + 1 < len(lines) and lines[i + 1].startswith('> '):
                i += 1
                quote_lines.append(lines[i][2:])
            blocks.append({'type': 'quote', 'content': '\n'.join(quote_lines)})
        
        # Regular paragraph (non-empty line)
        elif line.strip():
            blocks.append({'type': 'para', 'content': line.strip()})
        
        i += 1
    
    return blocks

def parse_table(table_lines):
    """Parse markdown table into headers and rows"""
    headers = []
    rows = []
    
    for i, line in enumerate(table_lines):
        cells = [c.strip() for c in line.split('|')]
        # Remove empty first/last cells from pipe boundaries
        cells = [c for c in cells if c]
        
        if i == 0:
            headers = cells
        elif i == 1 and all(c.replace('-', '').replace(':', '') == '' for c in cells):
            # Skip separator line
            continue
        else:
            if cells:
                rows.append(cells)
    
    return headers, rows

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
    text = re.sub(r'`([^`]+)`', r'\1', text)  # Inline code
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Links
    return text

def export_to_docx(blocks):
    """Export parsed blocks to DOCX"""
    doc = Document()
    
    for block in blocks:
        btype = block['type']
        
        if btype == 'h1':
            h = doc.add_heading(clean_markdown_text(block['content']), 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif btype == 'h2':
            doc.add_heading(clean_markdown_text(block['content']), 1)
        
        elif btype == 'h3':
            doc.add_heading(clean_markdown_text(block['content']), 2)
        
        elif btype == 'h4':
            doc.add_heading(clean_markdown_text(block['content']), 3)
        
        elif btype == 'para':
            text = clean_markdown_text(block['content'])
            if text:
                doc.add_paragraph(text)
        
        elif btype == 'quote':
            text = clean_markdown_text(block['content'])
            text = text.replace('[!IMPORTANT]', '⚠️ IMPORTANT:')
            text = text.replace('[!NOTE]', '📝 NOTE:')
            text = text.replace('[!WARNING]', '⚠️ WARNING:')
            text = text.replace('[!TIP]', '💡 TIP:')
            p = doc.add_paragraph()
            p.style = 'Quote'
            p.add_run(text)
        
        elif btype == 'code':
            p = doc.add_paragraph()
            run = p.add_run(block['content'])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        elif btype == 'table':
            headers, rows = parse_table(block['lines'])
            if headers and rows:
                # Ensure all rows have same number of columns
                ncols = len(headers)
                table = doc.add_table(rows=len(rows) + 1, cols=ncols)
                table.style = 'Table Grid'
                
                # Headers
                for j, h in enumerate(headers[:ncols]):
                    cell = table.rows[0].cells[j]
                    cell.text = clean_markdown_text(h)
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True
                
                # Rows
                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row[:ncols]):
                        if j < ncols:
                            table.rows[i + 1].cells[j].text = clean_markdown_text(cell_text)
                
                doc.add_paragraph()  # Space after table
        
        elif btype == 'image':
            img_path = fix_image_path(block['path'])
            full_path = os.path.join(BASE_DIR, img_path)
            if os.path.exists(full_path):
                doc.add_picture(full_path, width=Inches(6))
                # Caption
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(os.path.basename(img_path))
                run.italic = True
                run.font.size = Pt(10)
                print(f"✓ Added: {os.path.basename(img_path)}")
            else:
                print(f"✗ Missing: {os.path.basename(img_path)}")
        
        elif btype == 'hr':
            doc.add_paragraph('─' * 50)
    
    doc.save(OUTPUT_FILE)
    return len(blocks)

def main():
    print("📄 Exporting PROJECT_PAPER.md to DOCX...")
    print("=" * 50)
    
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        content = f.read()
    
    print(f"📝 Source: {len(content.split(chr(10)))} lines")
    
    blocks = parse_markdown(content)
    print(f"🔧 Parsed: {len(blocks)} blocks")
    
    export_to_docx(blocks)
    
    print(f"\n✅ Exported to: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
